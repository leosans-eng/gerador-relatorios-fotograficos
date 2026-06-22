"""Gera relatório fotográfico em Word a partir do modelo padrão."""

from __future__ import annotations

import math
import os
import re
import shutil
from copy import deepcopy
from pathlib import Path
from tkinter import filedialog, messagebox

from app_paths import app_dir, bundle_dir
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Emu
from docx.table import Table

_W_PPR_LOCAL_TAG = "pPr"
_W_T_TAG = qn("w:t")
_W_TR_TAG = qn("w:tr")
_HEADING1_TEXT = "REGISTROS FOTOGRÁFICOS "

APP_DIR = app_dir()
TEMPLATE_PATH = bundle_dir() / "modelos" / "relatorio_modelo.docx"
PHOTOS_PER_TABLE = 12
TABLE_COLS = 2
IMAGE_WIDTH = Emu(2667000)
IMAGE_HEIGHT = Emu(1917700)
_KEEP_BODY_PREFIX = 4  # espaço da capa, ÍNDICE, sumário e quebra após o índice


def resolve_template_path() -> Path:
    if TEMPLATE_PATH.exists():
        return TEMPLATE_PATH
    raise FileNotFoundError(
        "Modelo Word não encontrado. Coloque o arquivo em 'modelos/relatorio_modelo.docx'."
    )


def sanitize_filename_part(value: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*]', "", value.strip())
    return cleaned or "SEM-NOME"


def default_output_name(condominio_name: str, cidade: str = "", uf: str = "") -> str:
    condo = sanitize_filename_part(condominio_name).upper()
    city = sanitize_filename_part(cidade).upper() if cidade.strip() else "CIDADE"
    state = sanitize_filename_part(uf).upper() if uf.strip() else "UF"
    return f"2. RELATÓRIO FOTOGRÁFICO-ÁREA COMUM-COND. {condo}-{city}-{state}.docx"


def default_downloads_dir() -> Path:
    downloads = Path.home() / "Downloads"
    return downloads if downloads.exists() else Path.home()


def photos_without_anomaly(condominio_data: dict) -> list[tuple[str, int, str]]:
    missing = []
    for section in condominio_data.get("sections", []):
        section_name = section.get("name", "Sem nome")
        for photo in section.get("photos", []):
            if not str(photo.get("anomaly", "")).strip():
                missing.append(
                    (
                        section_name,
                        int(photo.get("order", 0)),
                        os.path.basename(str(photo.get("path", ""))),
                    )
                )
    return missing


def resolve_image_path(path_str: str) -> Path | None:
    if not path_str:
        return None
    path = Path(path_str)
    if path.is_file():
        return path
    relative = APP_DIR / path
    if relative.is_file():
        return relative
    return None


def _clear_paragraph(paragraph):
    element = paragraph._element
    for child in list(element):
        if child.tag.split("}")[-1] != _W_PPR_LOCAL_TAG:
            element.remove(child)


def _set_caption(paragraph, photo_number: int, anomaly: str):
    _clear_paragraph(paragraph)
    paragraph.add_run(f"Foto {photo_number} – {anomaly}")


def _set_image(paragraph, image_path: Path):
    _clear_paragraph(paragraph)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=IMAGE_WIDTH, height=IMAGE_HEIGHT)


def _clear_photo_cell(cell):
    for paragraph in cell.paragraphs:
        _clear_paragraph(paragraph)


def _fill_photo_cell(cell, image_path: Path, photo_number: int, anomaly: str):
    if len(cell.paragraphs) < 2:
        raise ValueError("Célula do modelo não possui estrutura esperada (imagem + legenda).")
    _set_image(cell.paragraphs[0], image_path)
    _set_caption(cell.paragraphs[1], photo_number, anomaly)


def _populate_table(table: Table, photos: list[dict]):
    needed_cells = len(photos)
    total_cells = len(table.rows) * TABLE_COLS
    for index in range(total_cells):
        row = index // TABLE_COLS
        col = index % TABLE_COLS
        cell = table.cell(row, col)
        if index >= needed_cells:
            _clear_photo_cell(cell)
            continue
        photo = photos[index]
        image_path = resolve_image_path(str(photo.get("path", "")))
        if image_path is None:
            raise FileNotFoundError(
                f"Imagem não encontrada para a foto {photo.get('order', '?')}: {photo.get('path', '')}"
            )
        _fill_photo_cell(
            cell,
            image_path,
            int(photo.get("order", index + 1)),
            str(photo.get("anomaly", "")).strip(),
        )


def _trim_table_rows(table_element, photo_count: int):
    needed_rows = max(1, math.ceil(photo_count / TABLE_COLS)) if photo_count else 0
    rows = table_element.findall(_W_TR_TAG)
    while len(rows) > needed_rows:
        table_element.remove(rows[-1])
        rows = table_element.findall(_W_TR_TAG)


def _remove_bookmarks(element):
    for tag in ("bookmarkStart", "bookmarkEnd"):
        for node in list(element.iter()):
            if node.tag.split("}")[-1] == tag:
                parent = node.getparent()
                if parent is not None:
                    parent.remove(node)


def _set_element_text(element, text: str):
    text_nodes = [node for node in element.iter() if node.tag == _W_T_TAG]
    if not text_nodes:
        run = element.makeelement(qn("w:r"), {})
        text_node = run.makeelement(_W_T_TAG, {})
        text_node.text = text
        run.append(text_node)
        element.append(run)
        return
    text_nodes[0].text = text
    for node in text_nodes[1:]:
        node.text = ""


def _get_num_id(template_paragraph) -> str | None:
    num_pr = template_paragraph._element.find(f".//{qn('w:numPr')}")
    if num_pr is None:
        return None
    num_id = num_pr.find(qn("w:numId"))
    if num_id is None:
        return None
    return num_id.get(qn("w:val"))


def _apply_heading_numbering(element, num_id: str, ilvl: str):
    p_pr = element.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = element.makeelement(qn("w:pPr"), {})
        element.insert(0, p_pr)
    existing = p_pr.find(qn("w:numPr"))
    if existing is not None:
        p_pr.remove(existing)
    num_pr = p_pr.makeelement(qn("w:numPr"), {})
    ilvl_el = num_pr.makeelement(qn("w:ilvl"), {qn("w:val"): ilvl})
    num_id_el = num_pr.makeelement(qn("w:numId"), {qn("w:val"): num_id})
    num_pr.append(ilvl_el)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)


def _clone_paragraph_element(
    template_paragraph,
    text: str | None = None,
    *,
    ilvl: str | None = None,
    num_id: str | None = None,
):
    element = deepcopy(template_paragraph._element)
    _remove_bookmarks(element)
    if text is not None:
        _set_element_text(element, text)
    if ilvl is not None and num_id:
        _apply_heading_numbering(element, num_id, ilvl)
    return element


def _insert_before_sectpr(doc, element):
    body = doc.element.body
    sect_pr = body[-1]
    body.insert(list(body).index(sect_pr), element)


def _insert_cover_placeholder(doc):
    """Parágrafo vazio no início do documento para o usuário colar a capa."""
    body = doc.element.body
    placeholder = doc.add_paragraph("")
    placeholder_element = placeholder._element
    body.remove(placeholder_element)
    body.insert(0, placeholder_element)


def _remove_generated_content(doc):
    body = doc.element.body
    children = list(body)
    if len(children) <= _KEEP_BODY_PREFIX + 1:
        return
    for child in children[_KEEP_BODY_PREFIX:-1]:
        body.remove(child)


def _paragraph_has_page_break(paragraph) -> bool:
    for br in paragraph._element.findall(f".//{qn('w:br')}"):
        if br.get(qn("w:type")) == "page":
            return True
    return False


def _is_heading_style(paragraph, level: int) -> bool:
    style_name = (paragraph.style.name or "").casefold()
    targets = {
        1: ("heading 1", "título 1", "titulo 1"),
        2: ("heading 2", "título 2", "titulo 2"),
    }
    return any(target in style_name for target in targets[level])


def _is_empty_normal_paragraph(paragraph) -> bool:
    if paragraph.text.strip():
        return False
    if _paragraph_has_page_break(paragraph):
        return False
    style_name = (paragraph.style.name or "Normal").casefold()
    return "normal" in style_name or style_name == ""


def _clone_empty_line_element(template_paragraph):
    element = deepcopy(template_paragraph._element)
    _remove_bookmarks(element)
    _set_element_text(element, "")
    return element


def _extract_template_parts(doc):
    heading1 = None
    heading2 = None
    page_break = None
    empty_line = None

    for paragraph in doc.paragraphs:
        if heading1 is None and _is_heading_style(paragraph, 1):
            heading1 = paragraph
            continue
        if heading2 is None and _is_heading_style(paragraph, 2):
            heading2 = paragraph
            continue
        if page_break is None and _paragraph_has_page_break(paragraph) and not paragraph.text.strip():
            page_break = paragraph
            continue
        if empty_line is None and _is_empty_normal_paragraph(paragraph):
            empty_line = paragraph

    if heading1 is None or heading2 is None or not doc.tables:
        raise ValueError(
            "O modelo Word não possui a estrutura esperada "
            "(título principal, título de bloco e tabela de fotos)."
        )

    if empty_line is None:
        empty_line = page_break

    return {
        "heading1": heading1,
        "heading2": heading2,
        "page_break": page_break,
        "empty_line": empty_line,
        "table": doc.tables[0]._tbl,
    }


def _format_section_heading(section_name: str) -> str:
    return section_name.strip().upper() or "BLOCO"


def _build_report_body(doc, condominio_data: dict):
    parts = _extract_template_parts(doc)
    template_table = parts["table"]
    _insert_cover_placeholder(doc)
    _remove_generated_content(doc)

    heading1 = parts["heading1"]
    heading_num_id = _get_num_id(heading1)
    _insert_before_sectpr(
        doc,
        _clone_paragraph_element(heading1, _HEADING1_TEXT),
    )

    sections = [
        section
        for section in condominio_data.get("sections", [])
        if section.get("photos")
    ]
    if not sections:
        return

    for section_index, section in enumerate(sections):
        photos = sorted(section.get("photos", []), key=lambda item: item.get("order", 0))
        _insert_before_sectpr(
            doc,
            _clone_paragraph_element(
                parts["heading2"],
                _format_section_heading(section.get("name", "Bloco")),
                ilvl="1",
                num_id=heading_num_id,
            ),
        )

        for chunk_start in range(0, len(photos), PHOTOS_PER_TABLE):
            chunk = photos[chunk_start : chunk_start + PHOTOS_PER_TABLE]
            new_table = deepcopy(template_table)
            _trim_table_rows(new_table, len(chunk))
            _insert_before_sectpr(doc, new_table)
            table = Table(new_table, doc)
            _populate_table(table, chunk)

        if section_index < len(sections) - 1 and parts["empty_line"] is not None:
            _insert_before_sectpr(
                doc,
                _clone_empty_line_element(parts["empty_line"]),
            )


def gerar_relatorio(
    condominio_name: str,
    condominio_data: dict,
    *,
    parent=None,
    output_path: Path | None = None,
    ask_save_path: bool = True,
    open_after_save: bool = True,
) -> Path:
    missing = photos_without_anomaly(condominio_data)
    if missing:
        raise ValueError("Existem fotos sem anomalia selecionada.")

    total_photos = sum(len(section.get("photos", [])) for section in condominio_data.get("sections", []))
    if total_photos == 0:
        raise ValueError("Não há fotos para gerar o relatório.")

    template_path = resolve_template_path()
    if output_path is None:
        suggested_name = default_output_name(
            condominio_name,
            str(condominio_data.get("cidade", "")),
            str(condominio_data.get("uf", "")),
        )
        initial_dir = default_downloads_dir()
        if ask_save_path:
            selected = filedialog.asksaveasfilename(
                parent=parent,
                title="Salvar relatório Word",
                defaultextension=".docx",
                filetypes=[("Documento Word", "*.docx")],
                initialdir=str(initial_dir),
                initialfile=suggested_name,
            )
            if not selected:
                raise RuntimeError("Salvamento cancelado pelo usuário.")
            output_path = Path(selected)
        else:
            output_path = initial_dir / suggested_name

    output_path = Path(output_path)
    if output_path.suffix.lower() != ".docx":
        output_path = output_path.with_suffix(".docx")

    shutil.copy2(template_path, output_path)
    doc = Document(str(output_path))
    _build_report_body(doc, condominio_data)
    doc.save(str(output_path))

    if open_after_save:
        os.startfile(str(output_path))

    return output_path


if __name__ == "__main__":
    import json
    import tkinter as tk

    data_file = APP_DIR / "condominios.json"
    if not data_file.exists():
        raise SystemExit("Arquivo condominios.json não encontrado.")

    with data_file.open("r", encoding="utf-8") as handle:
        data = json.load(handle)

    condominio_name = data.get("current_condominio") or ""
    if not condominio_name:
        raise SystemExit("Nenhum condomínio selecionado em condominios.json.")

    condominio_data = data["condominios"].get(condominio_name)
    if not condominio_data:
        raise SystemExit(f"Condomínio '{condominio_name}' não encontrado.")

    root = tk.Tk()
    root.withdraw()
    try:
        path = gerar_relatorio(condominio_name, condominio_data, parent=root)
        messagebox.showinfo("Relatório gerado", f"Arquivo salvo em:\n{path}", parent=root)
    except Exception as exc:
        messagebox.showerror("Erro ao gerar relatório", str(exc), parent=root)
    finally:
        root.destroy()
